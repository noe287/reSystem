#Nejat Onay Erkose - 2016
# nejatonay.erkose@airties.com

#!/usr/bin/env python

#importing docs to process microsoft docx formatted files                                                                                                                                                                                                                      
import docx, xlrd, xlwt, sys, os, subprocess
from docx.shared import Pt
from xlrd import open_workbook
from xlwt import easyxf
from xlutils.copy import copy
from subprocess import call

# Formats a table cell
def format_cell(cell, i, _size):
# replace table values
    for paragraph in cell[i].paragraphs:
        for run in paragraph.runs:
            run.font.size = Pt(_size)


# Formats a document line
def format_line(doc,_size):                                                                                                                                                                                                                                                    
    for run in doc.paragraphs[(len(doc.paragraphs) - 1)].runs:
        run.font.size = Pt(_size)

# Formats table entry
def format_table_entry(indx, _size, row, col, tables):
    for paragraph in tables[indx].cell(row, col).paragraphs:
        for run in paragraph.runs:
            run.font.size = Pt(_size)

def qa_release():
    doc = docx.Document(templatesDir+docname)
    tables = doc.tables
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(10)

    c=0
    r=0
    i=0
    # format_table_entry(0, 8, r, 1, tables)
    for row in tables[0].rows:                                                                                                                                                                                                                                                 
        for cell in row.cells:
            if c == 0 and r > 0:  # Parameter Column of Table - without header
                    if r == 3:
                            tables[0].cell(3,1).text='Not necessary'
                            continue
                    tables[0].cell(r, 1).text = path+'/'+targets[i] 
                    # format_table_entry(2, 10, r, 1, tables)
            c += 1
        c = 0
        i +=1
        if r == 0:
            i = 0
        
        if r == 3:
            i -= 1
            
        r += 1
        if i > 7:
            break

    #Add foss.tar.gz location 
    tables[1].cell(4,3).text=foss_target
    if not os.path.exists("./releaseoutput"):
            os.makedirs("./releaseoutput")
    doc.save("./releaseoutput/qa_release.docx")

def dvt_release():
    #rb=open_workbook("./templates/"+docname)
    #workbook = xlrd.open_workbook("./templates/"+docname)
    readbook = open_workbook(templatesDir+docname)
    readsheet = readbook.sheet_by_index(0)
    print "Doc read"
    writebook = copy(readbook)
    writesheet = writebook.get_sheet(0)
    #sheet.write(0,1,'=HYPERLINK("http://intranet.corp.airties.com/rnd/Shared%20Documents/_PROJECTS/Air%204920/Orange%204920/13.%20Firmwares/"+rver+"/AirTies_"+model+projname+"_FW_"+rver+".bin","Firmware")')
    writesheet.write(0, 1, model+projname)
    writesheet.write(1, 1, rver)
    writesheet.write(2, 1, xlwt.Formula('HYPERLINK("http://intranet.corp.airties.com/rnd/Shared%20Documents/_PROJECTS/Air%204920/Orange%204920/13.%20Firmwares/'+rver+'/AirTies_Air'+modproj+'_FW_'+rver+'.bin","Firmware")'))
    writesheet.write(5, 3, xlwt.Formula('HYPERLINK("'+path+targets[6]+'","MCP")'))
    writesheet.write(6, 3, xlwt.Formula('HYPERLINK("'+path+targets[7]+'","PSCS")'))
    writesheet.write(7, 3, xlwt.Formula('HYPERLINK("'+path+targets[2]+'","Mirror Image")'))
    writesheet.write(8, 3, xlwt.Formula('HYPERLINK("'+path+targets[0]+'","Release Notes")'))
    if not os.path.exists("./releaseoutput"):
            os.makedirs("./releaseoutput")
    writebook.save("./releaseoutput/dvt_release.xlsx")
    #workbook.save("./output/dvt_release.xlsx")

def md5Sum():
    i = 0
    for i in range(len(md5sumTargets)):
        call(['md5sum']+[md5sumTargets[i]])

#################################################################################################################
# Get the command line arguments #
#################################################################################################################
if len(sys.argv) < 4 :
    print "you must enter releasetype, templatedocname, projectname, root directory to the rnd, and fw version"                                                                                                                                                                                                
    exit()


templates={"dvt":"dvt_release_temp.xlsx","prmt":"prmt_release_temp.docx","qa":"qa_release_temp.docx"}
templatesDir = "/home/nejat/ContainSys/WorkDir/RELEASE/templates/"

#runs as below
#python release.py qa 4920FR-OR 1.25.4.2.335 http://intranet.corp.airties.com/rnd/Shared%20Documents/_PROJECTS/Air%204920/Orange%204920/13.%20Firmwares
#or releasepy.sh qa 4920FR-OR 1.25.4.2.513 http://intranet.corp.airties.com/rnd/Shared%20Documents/_PROJECTS/Air%204920/Orange%204920/13.%20Firmwares

releaseType = sys.argv[1]                  # dvt, qa, prmt
docname  = templates[releaseType]          # template name
# model = sys.argv[2]                      # 4920, 4820, 4830 etc...
# projname = sys.argv[3]                   # fr-or
productID = sys.argv[2]                    # 4920FR-OR
rver = sys.argv[3]                         # 1.25.4.2.259
remotedir = sys.argv[4]+'/'                # http://intranet.corp.airties.com/rnd/Shared%20Documents/Forms/AllItems.aspx?RootFolder=%2Frnd%2FShared%20Documents%2F%5FPROJECTS%2FAir%204920%2FOrange%204920%2F13%2E%20Firmwares/

foss_target=remotedir+rver+'/foss.tar.gz'
path = remotedir+rver
# modproj = model+projname
modproj = productID
# airmodproj ='AirTies_Air'+model+projname+'_FW'
airmodproj ='AirTies_Air'+modproj
airmodprojver = airmodproj+'_FW_'+rver

#Mirror image notation needs to be tweaked
print rver
print releaseType
print docname
print productID
print remotedir
ver=rver.split('.')
print ver[0]
print ver[1]
print ver[2]
print ver[3]
print ver[4]
mirver = ver[0]+'.'+ver[1]+'.'+ver[2]+'U.'+ver[3]+'.'+ver[4]
#################################################################################################################
# Define the target document names to be used in filling the templates                                          #
#################################################################################################################
targets = [
             airmodproj+"_Release_Notes_"+rver+".pdf",\
             airmodproj+"_FW_"+rver+".bin",\
             airmodproj+"_FW_"+mirver+".bin",\
             airmodproj+"_FW_"+rver+"_MUS.bin",\
             airmodproj+"_FW_"+rver+"_GENERIC_SIGN.bin",\
             airmodproj+"_FW_"+rver+"_Manufacturing_Upgrade.bin",\
             "Air"+modproj+"_MCP_v1.1.pdf",\
             "pscs-air"+modproj+"-1.1.pdf"
          ]

md5sumTargets = [
                    airmodprojver+".bin",
                    airmodprojver+"_GENERIC_SIGN.bin",
                    airmodprojver+"_Manufacturing_Upgrade.bin", 
                    airmodprojver+"_MUS.bin", 
                    airmodprojver+"_telnet_disabled_preinstall.bin", 
                    airmodproj+"_FW_"+mirver+".bin", 
                    airmodprojver+"_telnet_enabled_preinstall.bin",
                    "foss.tar.gz"
                ]   


#################################################################################################################
# use the docname and transform the ms docx format to be used in doc object                                     #
#################################################################################################################
if releaseType=="prmt" or releaseType=="qa":
    qa_release()
    md5Sum()
else:
    print "XSLS document read"
    dvt_release()

